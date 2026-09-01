import re
import unicodedata
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .models import RegistroTitulacion
from .services_excel import EXPORTAR_CAMPOS


def _texto(valor):
    if valor is None or valor == "":
        return ""
    if hasattr(valor, "strftime"):
        return valor.strftime("%d/%m/%Y")
    return str(valor)


def _nombre_seguro(registro, extension):
    texto = unicodedata.normalize("NFKD", registro.nombres_completos)
    texto = "".join(caracter for caracter in texto if not unicodedata.combining(caracter))
    texto = re.sub(r"[^A-Za-z0-9]+", "_", texto).strip("_") or "estudiante"
    return f"{texto}_{registro.cedula}.{extension}"


def _campos_documento(registro):
    return [
        (campo.verbose_name, _texto(getattr(registro, campo.name)))
        for campo in RegistroTitulacion._meta.fields
        if campo.name not in {"id", "fecha_creacion", "fecha_actualizacion"}
    ]


def crear_excel_estudiantes(registros):
    libro = Workbook()
    hoja = libro.active
    hoja.title = "Matriz de titulación"
    encabezados = [encabezado for encabezado, _ in EXPORTAR_CAMPOS]
    hoja.append(encabezados)

    for registro in registros:
        fila = []
        for encabezado, campo in EXPORTAR_CAMPOS:
            valor = getattr(registro, campo, "")
            fila.append(valor if valor is not None else "")
        hoja.append(fila)

    encabezado_fill = PatternFill("solid", fgColor="0B5D8A")
    encabezado_font = Font(color="FFFFFF", bold=True)
    for celda in hoja[1]:
        celda.fill = encabezado_fill
        celda.font = encabezado_font
        celda.alignment = Alignment(horizontal="center", vertical="center")
    hoja.auto_filter.ref = hoja.dimensions
    hoja.freeze_panes = "A2"
    hoja.row_dimensions[1].height = 24

    for indice, (_, campo) in enumerate(EXPORTAR_CAMPOS, start=1):
        valores = [hoja.cell(fila, indice).value for fila in range(1, hoja.max_row + 1)]
        ancho = min(max(len(str(valor or "")) for valor in valores) + 2, 42)
        hoja.column_dimensions[get_column_letter(indice)].width = max(ancho, 12)
        if campo == "fecha_grado":
            for fila in range(2, hoja.max_row + 1):
                hoja.cell(fila, indice).number_format = "dd/mm/yyyy"

    salida = BytesIO()
    libro.save(salida)
    return salida.getvalue()


def crear_pdf_estudiante(registro):
    salida = BytesIO()
    documento = SimpleDocTemplate(
        salida,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    estilos = getSampleStyleSheet()
    elementos = [
        Paragraph("PUCE TEC", estilos["Title"]),
        Paragraph("Información completa del estudiante", estilos["Heading2"]),
        Spacer(1, 8),
    ]
    filas = [["Campo", "Valor"]] + [[etiqueta, valor or "No registrado"] for etiqueta, valor in _campos_documento(registro)]
    tabla = Table(filas, colWidths=[65 * mm, 111 * mm], repeatRows=1)
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B5D8A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    elementos.append(tabla)
    documento.build(elementos)
    return salida.getvalue()


def crear_word_estudiante(registro):
    documento = Document()
    normal = documento.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    titulo = documento.add_heading("PUCE TEC", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = documento.add_heading("Información completa del estudiante", level=1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla = documento.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    tabla.rows[0].cells[0].text = "Campo"
    tabla.rows[0].cells[1].text = "Valor"
    for celda in tabla.rows[0].cells:
        celda.paragraphs[0].runs[0].bold = True
    for etiqueta, valor in _campos_documento(registro):
        celdas = tabla.add_row().cells
        celdas[0].text = etiqueta
        celdas[1].text = valor or "No registrado"
    salida = BytesIO()
    documento.save(salida)
    return salida.getvalue()


def exportar_masivo(registros, formato):
    registros = list(registros)
    if formato == "excel":
        return "estudiantes.xlsx", crear_excel_estudiantes(registros), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    extension = "pdf" if formato == "pdf" else "docx"
    crear_archivo = crear_pdf_estudiante if formato == "pdf" else crear_word_estudiante
    salida = BytesIO()
    with ZipFile(salida, "w", ZIP_DEFLATED) as archivo_zip:
        for registro in registros:
            archivo_zip.writestr(_nombre_seguro(registro, extension), crear_archivo(registro))
    return f"estudiantes_{extension}.zip", salida.getvalue(), "application/zip"
