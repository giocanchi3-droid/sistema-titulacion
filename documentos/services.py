from io import BytesIO
from textwrap import wrap

from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def mostrar(valor):
    if valor is None or valor == "":
        return "No registrado"

    return str(valor)


def obtener_nota_final(registro):
    if registro.modalidad_titulacion == "EXAMEN_COMPLEXIVO":
        return registro.nota_final2

    return registro.nota_final


def obtener_tribunal(registro):
    return [
        (
            "Primer miembro",
            registro.primer_miembro_tribunal,
            registro.primer_miembro_id_docente,
        ),
        (
            "Segundo miembro",
            registro.segundo_miembro_tribunal,
            registro.segundo_miembro_id_docente,
        ),
        (
            "Tercer miembro",
            registro.tercer_miembro_tribunal,
            registro.tercer_miembro_id_docente,
        ),
        (
            "Cuarto miembro",
            registro.cuarto_miembro_tribunal,
            registro.cuarto_miembro_id_docente,
        ),
    ]


def crear_word(acta):
    registro = acta.registro

    documento = Document()

    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(10)

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    texto_titulo = titulo.add_run(
        "PONTIFICIA UNIVERSIDAD CATÓLICA DEL ECUADOR"
    )
    texto_titulo.bold = True
    texto_titulo.font.size = Pt(13)

    subtitulo = documento.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    texto_subtitulo = subtitulo.add_run(
        "UNIDAD ACADÉMICA DE FORMACIÓN "
        "TÉCNICA Y TECNOLÓGICA – PUCE TEC"
    )
    texto_subtitulo.bold = True
    texto_subtitulo.font.size = Pt(11)

    encabezado = documento.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nombre_acta = encabezado.add_run(
        "\nACTA DE TITULACIÓN\n"
    )
    nombre_acta.bold = True
    nombre_acta.font.size = Pt(15)

    numero = documento.add_paragraph()
    numero.alignment = WD_ALIGN_PARAGRAPH.CENTER
    numero.add_run(
        f"N.º {acta.numero_acta}"
    ).bold = True

    documento.add_paragraph("")

    tabla = documento.add_table(
        rows=0,
        cols=2,
    )

    tabla.style = "Table Grid"

    informacion = [
        (
            "ID Banner",
            registro.id_banner,
        ),
        (
            "Nombres completos",
            registro.nombres_completos,
        ),
        (
            "Cédula",
            registro.cedula,
        ),
        (
            "Programa",
            registro.programa,
        ),
        (
            "Sede",
            registro.sede,
        ),
        (
            "Modalidad de titulación",
            registro.get_modalidad_titulacion_display(),
        ),
        (
            "Periodo de titulación",
            registro.periodo_titulacion_senescyt,
        ),
        (
            "Tema",
            registro.tema,
        ),
        (
            "Tutor",
            registro.nombres_completos_tutor,
        ),
        (
            "Identificación del tutor",
            registro.id_tutor,
        ),
        (
            "Nota final",
            obtener_nota_final(registro),
        ),
        (
            "Fecha de grado",
            (
                registro.fecha_grado.strftime("%d/%m/%Y")
                if registro.fecha_grado
                else "No registrada"
            ),
        ),
    ]

    for etiqueta, valor in informacion:
        celdas = tabla.add_row().cells

        celdas[0].text = etiqueta
        celdas[1].text = mostrar(valor)

        celdas[0].paragraphs[0].runs[0].bold = True

    documento.add_paragraph("")
    documento.add_heading(
        "Miembros del tribunal",
        level=2,
    )

    tabla_tribunal = documento.add_table(
        rows=1,
        cols=3,
    )

    tabla_tribunal.style = "Table Grid"

    encabezados = tabla_tribunal.rows[0].cells
    encabezados[0].text = "Cargo"
    encabezados[1].text = "Apellidos y nombres"
    encabezados[2].text = "Identificación"

    for celda in encabezados:
        celda.paragraphs[0].runs[0].bold = True

    for cargo, nombre, identificacion in obtener_tribunal(
        registro
    ):
        celdas = tabla_tribunal.add_row().cells

        celdas[0].text = cargo
        celdas[1].text = mostrar(nombre)
        celdas[2].text = mostrar(identificacion)

    documento.add_paragraph("")
    documento.add_heading(
        "Observaciones",
        level=2,
    )

    documento.add_paragraph(
        acta.observaciones
        or registro.observacion_secretaria
        or registro.observacion_puce_tec
        or "Sin observaciones."
    )

    documento.add_paragraph("\n\n")

    tabla_firmas = documento.add_table(
        rows=1,
        cols=2,
    )

    tabla_firmas.cells[0].text = (
        "____________________________\n"
        "Secretaría"
    )

    tabla_firmas.cells[1].text = (
        "____________________________\n"
        "Coordinación de carrera"
    )

    for celda in tabla_firmas.rows[0].cells:
        celda.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    salida = BytesIO()
    documento.save(salida)
    salida.seek(0)

    return salida


def crear_pdf(acta):
    registro = acta.registro

    salida = BytesIO()

    pdf = canvas.Canvas(
        salida,
        pagesize=A4,
    )

    ancho, alto = A4
    y = alto - 50

    def comprobar_pagina(espacio=25):
        nonlocal y

        if y < espacio:
            pdf.showPage()
            y = alto - 50

    def escribir_lineas(
        texto,
        x=55,
        ancho_texto=85,
        interlineado=14,
    ):
        nonlocal y

        for linea in wrap(
            mostrar(texto),
            width=ancho_texto,
        ):
            comprobar_pagina()

            pdf.drawString(
                x,
                y,
                linea,
            )

            y -= interlineado

    def campo(etiqueta, valor):
        nonlocal y

        comprobar_pagina(60)

        pdf.setFont(
            "Helvetica-Bold",
            10,
        )

        pdf.drawString(
            55,
            y,
            f"{etiqueta}:",
        )

        pdf.setFont(
            "Helvetica",
            10,
        )

        lineas = wrap(
            mostrar(valor),
            width=65,
        )

        if not lineas:
            lineas = ["No registrado"]

        pdf.drawString(
            190,
            y,
            lineas[0],
        )

        for linea in lineas[1:]:
            y -= 14

            pdf.drawString(
                190,
                y,
                linea,
            )

        y -= 20

    pdf.setFont(
        "Helvetica-Bold",
        13,
    )

    pdf.drawCentredString(
        ancho / 2,
        y,
        "PONTIFICIA UNIVERSIDAD CATÓLICA DEL ECUADOR",
    )

    y -= 20

    pdf.setFont(
        "Helvetica-Bold",
        9,
    )

    pdf.drawCentredString(
        ancho / 2,
        y,
        "UNIDAD ACADÉMICA DE FORMACIÓN "
        "TÉCNICA Y TECNOLÓGICA – PUCE TEC",
    )

    y -= 35

    pdf.setFont(
        "Helvetica-Bold",
        16,
    )

    pdf.drawCentredString(
        ancho / 2,
        y,
        "ACTA DE TITULACIÓN",
    )

    y -= 22

    pdf.setFont(
        "Helvetica-Bold",
        11,
    )

    pdf.drawCentredString(
        ancho / 2,
        y,
        f"N.º {acta.numero_acta}",
    )

    y -= 35

    campo(
        "ID Banner",
        registro.id_banner,
    )

    campo(
        "Estudiante",
        registro.nombres_completos,
    )

    campo(
        "Cédula",
        registro.cedula,
    )

    campo(
        "Programa",
        registro.programa,
    )

    campo(
        "Sede",
        registro.sede,
    )

    campo(
        "Modalidad",
        registro.get_modalidad_titulacion_display(),
    )

    campo(
        "Periodo",
        registro.periodo_titulacion_senescyt,
    )

    campo(
        "Tema",
        registro.tema,
    )

    campo(
        "Tutor",
        registro.nombres_completos_tutor,
    )

    campo(
        "Nota final",
        obtener_nota_final(registro),
    )

    campo(
        "Fecha de grado",
        (
            registro.fecha_grado.strftime("%d/%m/%Y")
            if registro.fecha_grado
            else "No registrada"
        ),
    )

    comprobar_pagina(150)

    pdf.setFont(
        "Helvetica-Bold",
        12,
    )

    pdf.drawString(
        55,
        y,
        "Miembros del tribunal",
    )

    y -= 22

    pdf.setFont(
        "Helvetica",
        9,
    )

    for cargo, nombre, identificacion in obtener_tribunal(
        registro
    ):
        texto = (
            f"{cargo}: {mostrar(nombre)} "
            f"- ID: {mostrar(identificacion)}"
        )

        escribir_lineas(
            texto,
            ancho_texto=95,
        )

        y -= 4

    comprobar_pagina(130)

    pdf.setFont(
        "Helvetica-Bold",
        12,
    )

    pdf.drawString(
        55,
        y,
        "Observaciones",
    )

    y -= 20

    pdf.setFont(
        "Helvetica",
        9,
    )

    observacion = (
        acta.observaciones
        or registro.observacion_secretaria
        or registro.observacion_puce_tec
        or "Sin observaciones."
    )

    escribir_lineas(
        observacion,
        ancho_texto=100,
    )

    comprobar_pagina(90)

    y -= 35

    pdf.line(
        70,
        y,
        250,
        y,
    )

    pdf.line(
        345,
        y,
        525,
        y,
    )

    pdf.setFont(
        "Helvetica",
        9,
    )

    pdf.drawCentredString(
        160,
        y - 15,
        "Firma de Secretaría",
    )

    pdf.drawCentredString(
        435,
        y - 15,
        "Firma de Coordinación",
    )

    pdf.save()
    salida.seek(0)

    return salida


def generar_archivos_acta(acta):
    word = crear_word(acta)
    pdf = crear_pdf(acta)

    nombre_base = (
        f"{acta.numero_acta}_"
        f"{acta.registro.cedula}"
    )

    if acta.archivo_word:
        acta.archivo_word.delete(
            save=False
        )

    if acta.archivo_pdf:
        acta.archivo_pdf.delete(
            save=False
        )

    acta.archivo_word.save(
        f"{nombre_base}.docx",
        ContentFile(
            word.getvalue()
        ),
        save=False,
    )

    acta.archivo_pdf.save(
        f"{nombre_base}.pdf",
        ContentFile(
            pdf.getvalue()
        ),
        save=False,
    )

    acta.estado = "GENERADA"
    acta.fecha_generacion = timezone.now()
    acta.save()

    return acta
