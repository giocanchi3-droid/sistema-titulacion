from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


DIAS = [
    "lunes", "martes", "miércoles", "jueves",
    "viernes", "sábado", "domingo"
]

MESES = [
    "", "enero", "febrero", "marzo", "abril",
    "mayo", "junio", "julio", "agosto",
    "septiembre", "octubre", "noviembre", "diciembre"
]


def valor(dato, vacio=""):
    if dato is None:
        return vacio

    texto = str(dato).strip()

    if not texto:
        return vacio

    try:
        numero = float(texto)

        if numero.is_integer():
            return str(int(numero))
    except Exception:
        pass

    return texto


def fecha_texto(fecha):
    if not fecha:
        fecha = date.today()

    return (
        f"{DIAS[fecha.weekday()]}, "
        f"{fecha.day} de {MESES[fecha.month]} de {fecha.year}"
    )


def es_complexivo(acta):
    tipo = str(
        getattr(acta, "tipo_acta", "")
        or ""
    ).lower()

    return "complex" in tipo


def miembro(registro, numero):
    campo = {
        1: "primer_miembro_tribunal",
        2: "segundo_miembro_tribunal",
        3: "tercer_miembro_tribunal",
        4: "cuarto_miembro_tribunal",
    }[numero]

    nombre = valor(
        getattr(registro, campo, "")
    )

    return nombre.upper() if nombre else "0"


def datos_acta(acta):
    r = acta.registro

    complexivo = es_complexivo(acta)

    if complexivo:
        titulo = "Registro de Defensa Examen Complexivo"

        presentacion = (
            "se presenta a la defensa de Examen Complexivo"
        )

        notas = [
            (
                "CALIFICACIÓN DEL\nEXAMEN TEÓRICO:",
                valor(
                    getattr(
                        r,
                        "examen_teorico_complexivo",
                        None
                    )
                )
            ),
            (
                "CALIFICACIÓN DEL\nEXAMEN PRÁCTICO:",
                (
                    valor(
                        getattr(
                            r,
                            "examen_teorico_practico",
                            None
                        )
                    )
                    + "/20"
                )
            ),
            (
                "SUMA TOTAL:",
                (
                    valor(
                        getattr(
                            r,
                            "nota_final2",
                            None
                        )
                    )
                    + "/50"
                )
            ),
        ]

        firmas = [
            miembro(r, 1),
            miembro(r, 2),
            miembro(r, 3),
        ]

    else:
        titulo = "Registro de Defensa oral"

        presentacion = (
            "se presenta a la exposición oral "
            "de su trabajo de titulación"
        )

        notas = [
            (
                "CALIFICACIÓN DEL TRABAJO ESCRITO:",
                valor(
                    getattr(
                        r,
                        "proyecto_escrito",
                        None
                    )
                )
            ),
            (
                "CALIFICACIÓN DE LA SUSTENTACIÓN ORAL:",
                (
                    valor(
                        getattr(
                            r,
                            "defensa_oral",
                            None
                        )
                    )
                    + "/20"
                )
            ),
            (
                "SUMA TOTAL:",
                (
                    valor(
                        getattr(
                            r,
                            "nota_final",
                            None
                        )
                    )
                    + "/50"
                )
            ),
        ]

        firmas = [
            miembro(r, 1),
            miembro(r, 2),
            miembro(r, 3),
            miembro(r, 4),
        ]

    return {
        "complexivo": complexivo,
        "titulo": titulo,

        "estudiante": valor(
            getattr(
                r,
                "nombres_completos",
                ""
            )
        ).upper(),

        "cedula": valor(
            getattr(
                r,
                "cedula",
                ""
            )
        ),

        "fecha": fecha_texto(
            getattr(
                r,
                "fecha_grado",
                None
            )
        ),

        "presentacion": presentacion,
        "notas": notas,
        "firmas": firmas,
    }


# ============================================================
# PDF
# ============================================================

def wrap(texto, fuente, tamano, ancho):
    resultado = []

    for bloque in str(texto).split("\n"):
        palabras = bloque.split()

        if not palabras:
            resultado.append("")
            continue

        linea = ""

        for palabra in palabras:
            nueva = (
                f"{linea} {palabra}".strip()
            )

            if stringWidth(
                nueva,
                fuente,
                tamano
            ) <= ancho:
                linea = nueva
            else:
                if linea:
                    resultado.append(linea)

                linea = palabra

        if linea:
            resultado.append(linea)

    return resultado


def texto_centrado(
    pdf,
    texto,
    x,
    y,
    ancho,
    fuente="Times-Roman",
    tamano=10,
    espacio=13
):
    pdf.setFont(fuente, tamano)

    for linea in wrap(
        texto,
        fuente,
        tamano,
        ancho
    ):
        pdf.drawCentredString(
            x + ancho / 2,
            y,
            linea
        )

        y -= espacio

    return y


def texto_izquierda(
    pdf,
    texto,
    x,
    y,
    ancho,
    fuente="Times-Roman",
    tamano=10,
    espacio=13
):
    pdf.setFont(fuente, tamano)

    for linea in wrap(
        texto,
        fuente,
        tamano,
        ancho
    ):
        pdf.drawString(
            x,
            y,
            linea
        )

        y -= espacio

    return y


def firma_pdf(
    pdf,
    x,
    y,
    nombre,
    mostrar_firma=True
):
    ancho = 158

    pdf.setLineWidth(1)

    pdf.line(
        x,
        y,
        x + ancho,
        y
    )

    y -= 12

    pdf.setFont(
        "Times-Roman",
        9
    )

    if mostrar_firma:
        pdf.drawCentredString(
            x + ancho / 2,
            y,
            "Firma"
        )

        y -= 14

    for linea in wrap(
        nombre,
        "Times-Roman",
        9,
        ancho
    ):
        pdf.drawCentredString(
            x + ancho / 2,
            y,
            linea
        )

        y -= 11


def crear_pdf(acta, logo):
    d = datos_acta(acta)

    buffer = BytesIO()

    pdf = canvas.Canvas(
        buffer,
        pagesize=A4
    )

    ancho, alto = A4

    # --------------------------------------------------------
    # ENCABEZADO
    # --------------------------------------------------------

    if Path(logo).exists():
        pdf.drawImage(
            str(logo),
            48,
            alto - 90,
            width=205,
            height=55,
            preserveAspectRatio=True,
            mask="auto"
        )

    pdf.setFont(
        "Helvetica",
        7.5
    )

    pdf.drawRightString(
        ancho - 50,
        alto - 49,
        "UNIDAD ACADÉMICA ESPECIALIZADA EN"
    )

    pdf.drawRightString(
        ancho - 50,
        alto - 60,
        "FORMACIÓN TÉCNICA Y TECNOLÓGICA"
    )

    pdf.drawRightString(
        ancho - 50,
        alto - 71,
        "PUCE TEC"
    )

    # --------------------------------------------------------
    # TITULO
    # --------------------------------------------------------

    pdf.setFont(
        "Times-Bold",
        12
    )

    pdf.drawCentredString(
        ancho / 2,
        alto - 120,
        d["titulo"]
    )

    # --------------------------------------------------------
    # PARRAFO PRINCIPAL
    # --------------------------------------------------------

    y = alto - 195

    intro = (
        "En la Ciudad de Quito el día, "
        f"{d['fecha']} , dando cumplimiento a las "
        "disposiciones en la normativa legal vigente, "
        "se deja constancia que el estudiante"
    )

    y = texto_centrado(
        pdf,
        intro,
        18,
        y,
        ancho - 36,
        tamano=10.5
    )

    pdf.setFont(
        "Times-Bold",
        10.5
    )

    pdf.drawCentredString(
        ancho / 2,
        y,
        d["estudiante"]
    )

    y -= 15

    texto_centrado(
        pdf,
        f"{d['cedula']}          {d['presentacion']}",
        20,
        y,
        ancho - 40,
        tamano=10.5
    )

    # --------------------------------------------------------
    # DELIBERACION
    # --------------------------------------------------------

    y -= 65

    y = texto_izquierda(
        pdf,
        (
            "Luego de la deliberación del tribunal "
            "el estudiante ha obtenido como resultado "
            "las siguientes calificaciones, en la "
            "materia de Integración Curricular"
        ),
        18,
        y,
        ancho - 36,
        tamano=10.5
    )

    y -= 17

    # --------------------------------------------------------
    # TABLA
    # --------------------------------------------------------

    if d["complexivo"]:
        tabla_x = 18
        tabla_ancho = 350
        col1 = 190

        alturas = [
            34,
            34,
            20
        ]

    else:
        tabla_x = 18
        tabla_ancho = ancho - 36

        col1 = tabla_ancho * 0.72

        alturas = [
            20,
            20,
            20
        ]

    col2 = tabla_ancho - col1

    for (
        etiqueta,
        nota
    ), altura_fila in zip(
        d["notas"],
        alturas
    ):
        pdf.setLineWidth(1.3)

        pdf.rect(
            tabla_x,
            y - altura_fila,
            col1,
            altura_fila
        )

        pdf.rect(
            tabla_x + col1,
            y - altura_fila,
            col2,
            altura_fila
        )

        lineas = wrap(
            etiqueta,
            "Times-Bold",
            10,
            col1 - 8
        )

        ty = y - 12

        pdf.setFont(
            "Times-Bold",
            10
        )

        for linea in lineas:
            pdf.drawString(
                tabla_x + 4,
                ty,
                linea
            )

            ty -= 11

        pdf.setFont(
            "Times-Roman",
            10.5
        )

        pdf.drawCentredString(
            tabla_x
            + col1
            + col2 / 2,
            y - altura_fila / 2 - 4,
            nota
        )

        y -= altura_fila

    # --------------------------------------------------------
    # AVISO
    # --------------------------------------------------------

    y -= 12

    texto_centrado(
        pdf,
        (
            "Se informa que esta nota y este evento "
            "no constituyen una ceremonia de graduación. "
            "De acuerdo con la normativa vigente de la "
            "PUCE, la Secretaría certificará el cumplimiento "
            "de todos los requisitos y notificará al "
            "estudiante correspondiente."
        ),
        22,
        y,
        ancho - 44,
        tamano=9.5,
        espacio=12
    )

    # --------------------------------------------------------
    # FIRMAS
    # --------------------------------------------------------

    if d["complexivo"]:
        posiciones = [
            (18, 265, False),
            (ancho - 178, 265, False),
            (18, 115, True),
        ]
    else:
        posiciones = [
            (18, 275, True),
            (ancho - 178, 275, True),
            (18, 120, True),
            (ancho - 178, 120, True),
        ]

    for posicion, nombre in zip(
        posiciones,
        d["firmas"]
    ):
        firma_pdf(
            pdf,
            posicion[0],
            posicion[1],
            nombre,
            posicion[2]
        )

    pdf.showPage()
    pdf.save()

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# WORD
# ============================================================

def sin_bordes(tabla):
    tbl_pr = tabla._tbl.tblPr

    borders = OxmlElement(
        "w:tblBorders"
    )

    for lado in [
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV"
    ]:
        nodo = OxmlElement(
            f"w:{lado}"
        )

        nodo.set(
            qn("w:val"),
            "nil"
        )

        borders.append(nodo)

    tbl_pr.append(borders)


def crear_word(acta, logo):
    d = datos_acta(acta)

    doc = Document()

    sec = doc.sections[0]

    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)

    sec.top_margin = Cm(1.1)
    sec.bottom_margin = Cm(1.1)

    sec.left_margin = Cm(1.25)
    sec.right_margin = Cm(1.25)

    normal = doc.styles["Normal"]

    normal.font.name = (
        "Times New Roman"
    )

    normal.font.size = Pt(10)

    # --------------------------------------------------------
    # ENCABEZADO
    # --------------------------------------------------------

    cabecera = doc.add_table(
        rows=1,
        cols=2
    )

    cabecera.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    sin_bordes(cabecera)

    izq = cabecera.cell(0, 0)
    der = cabecera.cell(0, 1)

    p = izq.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    if Path(logo).exists():
        p.add_run().add_picture(
            str(logo),
            width=Cm(6.0)
        )

    p = der.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    run = p.add_run(
        "UNIDAD ACADÉMICA ESPECIALIZADA EN\n"
        "FORMACIÓN TÉCNICA Y TECNOLÓGICA\n"
        "PUCE TEC"
    )

    run.font.name = "Arial"
    run.font.size = Pt(8)

    # --------------------------------------------------------
    # TITULO
    # --------------------------------------------------------

    p = doc.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(40)

    run = p.add_run(
        d["titulo"]
    )

    run.bold = True
    run.font.size = Pt(12)

    # --------------------------------------------------------
    # INTRODUCCION
    # --------------------------------------------------------

    p = doc.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.add_run(
        "En la Ciudad de Quito el día, "
        f"{d['fecha']} , dando cumplimiento a las\n"
        "disposiciones en la normativa legal vigente, "
        "se deja constancia que el estudiante"
    )

    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = p.add_run(
        d["estudiante"]
    )

    run.bold = True

    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.add_run(
        f"{d['cedula']}          "
        f"{d['presentacion']}"
    )

    p.paragraph_format.space_after = Pt(30)

    # --------------------------------------------------------
    # DELIBERACION
    # --------------------------------------------------------

    p = doc.add_paragraph(
        "Luego de la deliberación del tribunal "
        "el estudiante ha obtenido como resultado "
        "las siguientes\ncalificaciones, en la materia "
        "de Integración Curricular"
    )

    p.paragraph_format.space_after = Pt(10)

    # --------------------------------------------------------
    # TABLA NOTAS
    # --------------------------------------------------------

    tabla = doc.add_table(
        rows=0,
        cols=2
    )

    tabla.style = "Table Grid"

    tabla.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    tabla.autofit = False

    if d["complexivo"]:
        ancho1 = Cm(5.5)
        ancho2 = Cm(4.4)
    else:
        ancho1 = Cm(12.6)
        ancho2 = Cm(4.6)

    for etiqueta, nota in d["notas"]:
        celdas = tabla.add_row().cells

        celdas[0].width = ancho1
        celdas[1].width = ancho2

        celdas[0].vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        celdas[1].vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        p1 = celdas[0].paragraphs[0]

        p1.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = p1.add_run(
            etiqueta
        )

        run.bold = True
        run.font.size = Pt(10)

        p2 = celdas[1].paragraphs[0]

        p2.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = p2.add_run(
            nota
        )

        run.font.size = Pt(10)

    # --------------------------------------------------------
    # AVISO
    # --------------------------------------------------------

    p = doc.add_paragraph(
        "Se informa que esta nota y este evento "
        "no constituyen una ceremonia de graduación. "
        "De acuerdo con\nla normativa vigente de la "
        "PUCE, la Secretaría certificará el cumplimiento "
        "de todos los requisitos y\nnotificará al "
        "estudiante correspondiente."
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = Pt(52)

    # --------------------------------------------------------
    # FIRMAS
    # --------------------------------------------------------

    firmas = doc.add_table(
        rows=2,
        cols=2
    )

    firmas.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    sin_bordes(firmas)

    indice = 0

    for fila in range(2):
        for col in range(2):
            celda = firmas.cell(
                fila,
                col
            )

            p = celda.paragraphs[0]

            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            if fila == 1:
                p.paragraph_format.space_before = Pt(55)

            if indice < len(
                d["firmas"]
            ):
                p.add_run(
                    "____________________________\n"
                )

                mostrar_firma = (
                    not d["complexivo"]
                    or indice == 2
                )

                if mostrar_firma:
                    p.add_run(
                        "Firma\n"
                    )

                p.add_run(
                    d["firmas"][indice]
                )

            indice += 1

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()
