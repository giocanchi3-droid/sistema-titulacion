from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


DIAS = [
    "lunes",
    "martes",
    "miércoles",
    "jueves",
    "viernes",
    "sábado",
    "domingo",
]

MESES = [
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]


def _texto(valor):
    if valor is None:
        return "—"

    texto = str(valor).strip()

    if not texto:
        return "—"

    if texto.endswith(".0"):
        try:
            return str(int(float(texto)))
        except ValueError:
            pass

    return texto


def _nota(valor, maximo=None):
    valor_texto = _texto(valor)

    if maximo is None:
        return valor_texto

    if valor_texto == "—":
        return f"—/{maximo}"

    return f"{valor_texto}/{maximo}"


def _es_complexivo(acta):
    tipo = str(
        getattr(acta, "tipo_acta", "")
    ).lower()

    return "complex" in tipo


def _fecha_acta(registro):
    fecha = getattr(
        registro,
        "fecha_grado",
        None,
    )

    if not fecha:
        fecha = date.today()

    dia = DIAS[fecha.weekday()]
    mes = MESES[fecha.month]

    return (
        f"{dia}, {fecha.day} de "
        f"{mes} de {fecha.year}"
    )


def _nombre_tribunal(registro, campo):
    valor = getattr(
        registro,
        campo,
        "",
    )

    valor = str(valor or "").strip()

    if valor:
        return valor.upper()

    return "MIEMBRO DEL TRIBUNAL"


def contexto_acta(acta):
    registro = acta.registro

    complexivo = _es_complexivo(
        acta
    )

    if complexivo:

        titulo = (
            "Registro de Defensa "
            "Examen Complexivo"
        )

        frase = (
            "se presenta a la defensa "
            "de Examen Complexivo"
        )

        calificaciones = [
            {
                "label":
                    "CALIFICACIÓN DEL "
                    "EXAMEN TEÓRICO:",
                "valor": _texto(
                    getattr(
                        registro,
                        "examen_teorico_complexivo",
                        None,
                    )
                ),
            },
            {
                "label":
                    "CALIFICACIÓN DEL "
                    "EXAMEN PRÁCTICO:",
                "valor": _nota(
                    getattr(
                        registro,
                        "examen_teorico_practico",
                        None,
                    ),
                    20,
                ),
            },
            {
                "label": "SUMA TOTAL:",
                "valor": _nota(
                    getattr(
                        registro,
                        "nota_final2",
                        None,
                    ),
                    50,
                ),
            },
        ]

        firmas = [
            _nombre_tribunal(
                registro,
                "primer_miembro_tribunal",
            ),
            _nombre_tribunal(
                registro,
                "segundo_miembro_tribunal",
            ),
            _nombre_tribunal(
                registro,
                "tercer_miembro_tribunal",
            ),
        ]

    else:

        titulo = "Registro de Defensa oral"

        frase = (
            "se presenta a la exposición "
            "oral de su trabajo de titulación"
        )

        calificaciones = [
            {
                "label":
                    "CALIFICACIÓN DEL "
                    "TRABAJO ESCRITO:",
                "valor": _texto(
                    getattr(
                        registro,
                        "proyecto_escrito",
                        None,
                    )
                ),
            },
            {
                "label":
                    "CALIFICACIÓN DE LA "
                    "SUSTENTACIÓN ORAL:",
                "valor": _nota(
                    getattr(
                        registro,
                        "defensa_oral",
                        None,
                    ),
                    20,
                ),
            },
            {
                "label": "SUMA TOTAL:",
                "valor": _nota(
                    getattr(
                        registro,
                        "nota_final",
                        None,
                    ),
                    50,
                ),
            },
        ]

        firmas = [
            _nombre_tribunal(
                registro,
                "primer_miembro_tribunal",
            ),
            _nombre_tribunal(
                registro,
                "segundo_miembro_tribunal",
            ),
            _nombre_tribunal(
                registro,
                "tercer_miembro_tribunal",
            ),
            _nombre_tribunal(
                registro,
                "cuarto_miembro_tribunal",
            ),
        ]

    return {
        "acta": acta,
        "registro": registro,
        "complexivo": complexivo,
        "titulo_documento": titulo,
        "fecha_texto": _fecha_acta(
            registro
        ),
        "nombre_estudiante": str(
            getattr(
                registro,
                "nombres_completos",
                "",
            )
            or ""
        ).upper(),
        "cedula": _texto(
            getattr(
                registro,
                "cedula",
                "",
            )
        ),
        "frase_presentacion": frase,
        "calificaciones": calificaciones,
        "firmas": firmas,
    }


def _fuente_parrafo(
    parrafo,
    size=10,
    bold=False,
):
    for run in parrafo.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def generar_word_oficial(
    acta,
    logo_path,
):
    datos = contexto_acta(acta)

    documento = Document()

    seccion = documento.sections[0]

    seccion.page_width = Cm(21)
    seccion.page_height = Cm(29.7)

    seccion.top_margin = Cm(1.3)
    seccion.bottom_margin = Cm(1.3)
    seccion.left_margin = Cm(1.4)
    seccion.right_margin = Cm(1.4)

    normal = documento.styles["Normal"]

    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    # --------------------------------------------------------
    # CABECERA
    # --------------------------------------------------------

    cabecera = documento.add_table(
        rows=1,
        cols=2,
    )

    cabecera.autofit = False

    izquierda = cabecera.cell(0, 0)
    derecha = cabecera.cell(0, 1)

    izquierda.width = Cm(8.5)
    derecha.width = Cm(9.5)

    izquierda.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    derecha.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    p_logo = izquierda.paragraphs[0]

    p_logo.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    logo_path = Path(logo_path)

    if logo_path.exists():
        p_logo.add_run().add_picture(
            str(logo_path),
            width=Cm(7.0),
        )

    p_unidad = derecha.paragraphs[0]

    p_unidad.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    p_unidad.add_run(
        "UNIDAD ACADÉMICA ESPECIALIZADA EN\n"
        "FORMACIÓN TÉCNICA Y TECNOLÓGICA\n"
        "PUCE TEC"
    )

    _fuente_parrafo(
        p_unidad,
        size=8,
    )

    # --------------------------------------------------------
    # TITULO
    # --------------------------------------------------------

    titulo = documento.add_paragraph()

    titulo.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    titulo.paragraph_format.space_before = (
        Pt(8)
    )

    titulo.paragraph_format.space_after = (
        Pt(28)
    )

    run = titulo.add_run(
        datos["titulo_documento"]
    )

    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # --------------------------------------------------------
    # PARRAFO PRINCIPAL
    # --------------------------------------------------------

    p = documento.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    p.paragraph_format.space_after = Pt(16)

    p.add_run(
        "En la Ciudad de Quito el día, "
    )

    fecha_run = p.add_run(
        datos["fecha_texto"]
    )

    fecha_run.bold = False

    p.add_run(
        ", dando cumplimiento a las "
        "disposiciones en la normativa legal "
        "vigente, se deja constancia que el "
        "estudiante "
    )

    nombre_run = p.add_run(
        datos["nombre_estudiante"]
    )

    nombre_run.bold = True

    p.add_run(
        " con número de cédula "
    )

    cedula_run = p.add_run(
        datos["cedula"]
    )

    cedula_run.bold = False

    p.add_run(
        " "
        + datos["frase_presentacion"]
        + "."
    )

    _fuente_parrafo(
        p,
        size=10,
    )

    nombre_run.bold = True

    # --------------------------------------------------------
    # DELIBERACION
    # --------------------------------------------------------

    p = documento.add_paragraph(
        "Luego de la deliberación del tribunal "
        "el estudiante ha obtenido como resultado "
        "las siguientes calificaciones, en la "
        "materia de Integración Curricular"
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    p.paragraph_format.space_after = Pt(10)

    _fuente_parrafo(
        p,
        size=10,
    )

    # --------------------------------------------------------
    # TABLA DE CALIFICACIONES
    # --------------------------------------------------------

    tabla = documento.add_table(
        rows=0,
        cols=2,
    )

    tabla.style = "Table Grid"

    tabla.autofit = False

    for fila in datos["calificaciones"]:

        celdas = tabla.add_row().cells

        celdas[0].width = Cm(12.5)
        celdas[1].width = Cm(4.5)

        p1 = celdas[0].paragraphs[0]
        p2 = celdas[1].paragraphs[0]

        p1.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        p2.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        r1 = p1.add_run(
            fila["label"]
        )

        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(10)

        r2 = p2.add_run(
            fila["valor"]
        )

        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)

    # --------------------------------------------------------
    # AVISO
    # --------------------------------------------------------

    aviso = documento.add_paragraph(
        "Se informa que esta nota y este evento "
        "no constituyen una ceremonia de graduación. "
        "De acuerdo con la normativa vigente de la "
        "PUCE, la Secretaría certificará el "
        "cumplimiento de todos los requisitos y "
        "notificará al estudiante correspondiente."
    )

    aviso.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    aviso.paragraph_format.space_after = (
        Pt(45)
    )

    _fuente_parrafo(
        aviso,
        size=9,
    )

    # --------------------------------------------------------
    # FIRMAS
    # --------------------------------------------------------

    numero_firmas = len(
        datos["firmas"]
    )

    filas = 2
    columnas = 2

    firmas_tabla = documento.add_table(
        rows=filas,
        cols=columnas,
    )

    indice = 0

    for fila in range(filas):
        for columna in range(columnas):

            celda = firmas_tabla.cell(
                fila,
                columna,
            )

            p = celda.paragraphs[0]

            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            if indice < numero_firmas:

                p.add_run(
                    "\n__________________________\n"
                )

                firma_run = p.add_run(
                    "Firma\n"
                )

                firma_run.font.size = Pt(9)

                nombre_run = p.add_run(
                    datos["firmas"][indice]
                )

                nombre_run.font.name = (
                    "Times New Roman"
                )

                nombre_run.font.size = Pt(9)

            indice += 1

            if fila == 0:
                p.paragraph_format.space_after = (
                    Pt(45)
                )

    memoria = BytesIO()

    documento.save(
        memoria
    )

    memoria.seek(0)

    return memoria.getvalue()


# ============================================================
# PDF
# ============================================================

def _wrap(
    texto,
    fuente,
    tamano,
    ancho,
):
    palabras = texto.split()

    lineas = []
    actual = ""

    for palabra in palabras:

        propuesta = (
            f"{actual} {palabra}".strip()
        )

        if (
            stringWidth(
                propuesta,
                fuente,
                tamano,
            )
            <= ancho
        ):
            actual = propuesta
        else:
            if actual:
                lineas.append(actual)

            actual = palabra

    if actual:
        lineas.append(actual)

    return lineas


def _dibujar_parrafo(
    pdf,
    texto,
    x,
    y,
    ancho,
    fuente="Times-Roman",
    tamano=10,
    interlineado=13,
    centrado=False,
):
    lineas = _wrap(
        texto,
        fuente,
        tamano,
        ancho,
    )

    pdf.setFont(
        fuente,
        tamano,
    )

    for linea in lineas:

        if centrado:

            centro = x + (
                ancho / 2
            )

            pdf.drawCentredString(
                centro,
                y,
                linea,
            )

        else:

            pdf.drawString(
                x,
                y,
                linea,
            )

        y -= interlineado

    return y


def _dibujar_firma(
    pdf,
    x,
    y,
    nombre,
    ancho=150,
):
    pdf.setLineWidth(1)

    pdf.line(
        x,
        y,
        x + ancho,
        y,
    )

    pdf.setFont(
        "Times-Roman",
        9,
    )

    pdf.drawCentredString(
        x + ancho / 2,
        y - 12,
        "Firma",
    )

    lineas = _wrap(
        nombre,
        "Times-Roman",
        9,
        ancho,
    )

    pos_y = y - 26

    for linea in lineas:

        pdf.drawCentredString(
            x + ancho / 2,
            pos_y,
            linea,
        )

        pos_y -= 11


def generar_pdf_oficial(
    acta,
    logo_path,
):
    datos = contexto_acta(acta)

    memoria = BytesIO()

    pdf = canvas.Canvas(
        memoria,
        pagesize=A4,
    )

    ancho_pagina, alto_pagina = A4

    # --------------------------------------------------------
    # LOGO PUCE TEC
    # --------------------------------------------------------

    logo_path = Path(logo_path)

    if logo_path.exists():

        pdf.drawImage(
            str(logo_path),
            52,
            alto_pagina - 105,
            width=195,
            height=62,
            preserveAspectRatio=True,
            anchor="c",
            mask="auto",
        )

    # --------------------------------------------------------
    # UNIDAD
    # --------------------------------------------------------

    pdf.setFont(
        "Helvetica",
        7.7,
    )

    x_unidad = ancho_pagina - 55

    pdf.drawRightString(
        x_unidad,
        alto_pagina - 62,
        "UNIDAD ACADÉMICA ESPECIALIZADA EN",
    )

    pdf.drawRightString(
        x_unidad,
        alto_pagina - 74,
        "FORMACIÓN TÉCNICA Y TECNOLÓGICA",
    )

    pdf.drawRightString(
        x_unidad,
        alto_pagina - 86,
        "PUCE TEC",
    )

    # --------------------------------------------------------
    # TITULO
    # --------------------------------------------------------

    pdf.setFont(
        "Times-Bold",
        12,
    )

    pdf.drawCentredString(
        ancho_pagina / 2,
        alto_pagina - 135,
        datos["titulo_documento"],
    )

    # --------------------------------------------------------
    # TEXTO PRINCIPAL
    # --------------------------------------------------------

    y = alto_pagina - 205

    texto_inicio = (
        "En la Ciudad de Quito el día, "
        f"{datos['fecha_texto']}, "
        "dando cumplimiento a las disposiciones "
        "en la normativa legal vigente, se deja "
        "constancia que el estudiante"
    )

    y = _dibujar_parrafo(
        pdf,
        texto_inicio,
        55,
        y,
        ancho_pagina - 110,
        fuente="Times-Roman",
        tamano=10,
        interlineado=13,
        centrado=True,
    )

    y -= 2

    pdf.setFont(
        "Times-Bold",
        10,
    )

    pdf.drawCentredString(
        ancho_pagina / 2,
        y,
        datos["nombre_estudiante"],
    )

    y -= 15

    texto_estudiante = (
        f"{datos['cedula']} "
        f"{datos['frase_presentacion']}"
    )

    y = _dibujar_parrafo(
        pdf,
        texto_estudiante,
        55,
        y,
        ancho_pagina - 110,
        fuente="Times-Roman",
        tamano=10,
        interlineado=13,
        centrado=True,
    )

    y -= 35

    deliberacion = (
        "Luego de la deliberación del tribunal "
        "el estudiante ha obtenido como resultado "
        "las siguientes calificaciones, en la "
        "materia de Integración Curricular"
    )

    y = _dibujar_parrafo(
        pdf,
        deliberacion,
        55,
        y,
        ancho_pagina - 110,
        fuente="Times-Roman",
        tamano=10,
        interlineado=13,
    )

    y -= 16

    # --------------------------------------------------------
    # TABLA
    # --------------------------------------------------------

    if datos["complexivo"]:
        tabla_ancho = 360
    else:
        tabla_ancho = 470

    tabla_x = 55

    columna_1 = tabla_ancho * 0.70
    columna_2 = tabla_ancho * 0.30

    fila_altura = 28

    for fila in datos["calificaciones"]:

        pdf.rect(
            tabla_x,
            y - fila_altura,
            columna_1,
            fila_altura,
        )

        pdf.rect(
            tabla_x + columna_1,
            y - fila_altura,
            columna_2,
            fila_altura,
        )

        pdf.setFont(
            "Times-Bold",
            9,
        )

        lineas_label = _wrap(
            fila["label"],
            "Times-Bold",
            9,
            columna_1 - 10,
        )

        ly = y - 11

        for linea in lineas_label:

            pdf.drawString(
                tabla_x + 5,
                ly,
                linea,
            )

            ly -= 10

        pdf.setFont(
            "Times-Roman",
            10,
        )

        pdf.drawCentredString(
            tabla_x
            + columna_1
            + columna_2 / 2,
            y - 17,
            fila["valor"],
        )

        y -= fila_altura

    # --------------------------------------------------------
    # AVISO
    # --------------------------------------------------------

    y -= 12

    aviso = (
        "Se informa que esta nota y este evento "
        "no constituyen una ceremonia de graduación. "
        "De acuerdo con la normativa vigente de la "
        "PUCE, la Secretaría certificará el cumplimiento "
        "de todos los requisitos y notificará al "
        "estudiante correspondiente."
    )

    y = _dibujar_parrafo(
        pdf,
        aviso,
        65,
        y,
        ancho_pagina - 130,
        fuente="Times-Roman",
        tamano=9,
        interlineado=12,
        centrado=True,
    )

    # --------------------------------------------------------
    # FIRMAS
    # --------------------------------------------------------

    if datos["complexivo"]:

        posiciones = [
            (55, 265),
            (365, 265),
            (55, 115),
        ]

    else:

        posiciones = [
            (55, 270),
            (365, 270),
            (55, 120),
            (365, 120),
        ]

    for posicion, nombre in zip(
        posiciones,
        datos["firmas"],
    ):

        _dibujar_firma(
            pdf,
            posicion[0],
            posicion[1],
            nombre,
            ancho=150,
        )

    pdf.showPage()
    pdf.save()

    memoria.seek(0)

    return memoria.getvalue()
