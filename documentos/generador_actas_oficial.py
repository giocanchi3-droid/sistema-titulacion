from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


DIAS = [
    "lunes",
    "martes",
    "miércoles",
    "jueves",
    "viernes",
    "sábado",
    "domingo",
]

MESES = [
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]


def texto(valor, defecto="—"):
    if valor is None:
        return defecto

    resultado = str(valor).strip()

    if not resultado:
        return defecto

    try:
        numero = float(resultado)

        if numero.is_integer():
            return str(int(numero))

    except (ValueError, TypeError):
        pass

    return resultado


def fecha_espanol(fecha):
    if not fecha:
        fecha = date.today()

    return (
        f"{DIAS[fecha.weekday()]}, "
        f"{fecha.day} de "
        f"{MESES[fecha.month]} de "
        f"{fecha.year}"
    )


def es_complexivo(acta):
    tipo = str(
        getattr(acta, "tipo_acta", "")
    ).lower()

    return "complex" in tipo


def nombre_tribunal(registro, campo):
    valor = getattr(
        registro,
        campo,
        "",
    )

    valor = str(
        valor or ""
    ).strip()

    if not valor:
        return "—"

    return valor.upper()


def contexto_acta(acta):
    registro = acta.registro

    complexivo = es_complexivo(
        acta
    )

    if complexivo:
        titulo = (
            "Registro de Defensa "
            "Examen Complexivo"
        )

        presentacion = (
            "se presenta a la defensa "
            "de Examen Complexivo"
        )

        calificaciones = [
            (
                "CALIFICACIÓN DEL\n"
                "EXAMEN TEÓRICO:",
                texto(
                    getattr(
                        registro,
                        "examen_teorico_complexivo",
                        None,
                    )
                ),
            ),
            (
                "CALIFICACIÓN DEL\n"
                "EXAMEN PRÁCTICO:",
                f"{texto(getattr(registro, 'examen_teorico_practico', None), '')}/20",
            ),
            (
                "SUMA TOTAL:",
                f"{texto(getattr(registro, 'nota_final2', None), '')}/50",
            ),
        ]

        firmas = [
            nombre_tribunal(
                registro,
                "primer_miembro_tribunal",
            ),
            nombre_tribunal(
                registro,
                "segundo_miembro_tribunal",
            ),
            nombre_tribunal(
                registro,
                "tercer_miembro_tribunal",
            ),
        ]

    else:
        titulo = "Registro de Defensa oral"

        presentacion = (
            "se presenta a la exposición oral "
            "de su trabajo de titulación"
        )

        calificaciones = [
            (
                "CALIFICACIÓN DEL TRABAJO ESCRITO:",
                texto(
                    getattr(
                        registro,
                        "proyecto_escrito",
                        None,
                    )
                ),
            ),
            (
                "CALIFICACIÓN DE LA SUSTENTACIÓN ORAL:",
                f"{texto(getattr(registro, 'defensa_oral', None), '')}/20",
            ),
            (
                "SUMA TOTAL:",
                f"{texto(getattr(registro, 'nota_final', None), '')}/50",
            ),
        ]

        firmas = [
            nombre_tribunal(
                registro,
                "primer_miembro_tribunal",
            ),
            nombre_tribunal(
                registro,
                "segundo_miembro_tribunal",
            ),
            nombre_tribunal(
                registro,
                "tercer_miembro_tribunal",
            ),
            nombre_tribunal(
                registro,
                "cuarto_miembro_tribunal",
            ),
        ]

    return {
        "registro": registro,
        "complexivo": complexivo,
        "titulo": titulo,
        "presentacion": presentacion,
        "fecha": fecha_espanol(
            getattr(
                registro,
                "fecha_grado",
                None,
            )
        ),
        "estudiante": str(
            getattr(
                registro,
                "nombres_completos",
                "",
            )
            or ""
        ).upper(),
        "cedula": texto(
            getattr(
                registro,
                "cedula",
                "",
            )
        ),
        "calificaciones": calificaciones,
        "firmas": firmas,
    }


# ==========================================================
# PDF
# ==========================================================

def dividir_lineas(
    contenido,
    fuente,
    tamano,
    ancho,
):
    resultado = []

    for parrafo in str(
        contenido
    ).split("\n"):

        palabras = parrafo.split()

        if not palabras:
            resultado.append("")
            continue

        linea = ""

        for palabra in palabras:
            prueba = (
                f"{linea} {palabra}".strip()
            )

            if (
                stringWidth(
                    prueba,
                    fuente,
                    tamano,
                )
                <= ancho
            ):
                linea = prueba

            else:
                if linea:
                    resultado.append(
                        linea
                    )

                linea = palabra

        if linea:
            resultado.append(linea)

    return resultado


def parrafo_centrado(
    pdf,
    contenido,
    x,
    y,
    ancho,
    fuente="Times-Roman",
    tamano=10,
    interlineado=12,
):
    lineas = dividir_lineas(
        contenido,
        fuente,
        tamano,
        ancho,
    )

    pdf.setFont(
        fuente,
        tamano,
    )

    for linea in lineas:
        pdf.drawCentredString(
            x + ancho / 2,
            y,
            linea,
        )

        y -= interlineado

    return y


def parrafo_izquierda(
    pdf,
    contenido,
    x,
    y,
    ancho,
    fuente="Times-Roman",
    tamano=10,
    interlineado=12,
):
    lineas = dividir_lineas(
        contenido,
        fuente,
        tamano,
        ancho,
    )

    pdf.setFont(
        fuente,
        tamano,
    )

    for linea in lineas:
        pdf.drawString(
            x,
            y,
            linea,
        )

        y -= interlineado

    return y


def dibujar_firma(
    pdf,
    x,
    y,
    nombre,
    mostrar_firma=True,
):
    ancho = 160

    pdf.setLineWidth(1)

    pdf.line(
        x,
        y,
        x + ancho,
        y,
    )

    posicion = y - 12

    pdf.setFont(
        "Times-Roman",
        9.5,
    )

    if mostrar_firma:
        pdf.drawCentredString(
            x + ancho / 2,
            posicion,
            "Firma",
        )

        posicion -= 14

    lineas = dividir_lineas(
        nombre,
        "Times-Roman",
        9.5,
        ancho,
    )

    for linea in lineas:
        pdf.drawCentredString(
            x + ancho / 2,
            posicion,
            linea,
        )

        posicion -= 11


def generar_pdf_oficial(
    acta,
    logo_path,
):
    datos = contexto_acta(
        acta
    )

    memoria = BytesIO()

    pdf = canvas.Canvas(
        memoria,
        pagesize=A4,
    )

    ancho_pagina, alto_pagina = A4

    logo_path = Path(
        logo_path
    )

    if logo_path.exists():
        pdf.drawImage(
            str(logo_path),
            48,
            alto_pagina - 92,
            width=205,
            height=58,
            preserveAspectRatio=True,
            anchor="w",
            mask="auto",
        )

    # Encabezado derecho

    pdf.setFont(
        "Helvetica",
        7.7,
    )

    derecha = ancho_pagina - 50

    pdf.drawRightString(
        derecha,
        alto_pagina - 51,
        "UNIDAD ACADÉMICA ESPECIALIZADA EN",
    )

    pdf.drawRightString(
        derecha,
        alto_pagina - 62,
        "FORMACIÓN TÉCNICA Y TECNOLÓGICA",
    )

    pdf.drawRightString(
        derecha,
        alto_pagina - 73,
        "PUCE TEC",
    )

    # Título

    pdf.setFont(
        "Times-Bold",
        12,
    )

    pdf.drawCentredString(
        ancho_pagina / 2,
        alto_pagina - 118,
        datos["titulo"],
    )

    y = alto_pagina - 190

    introduccion = (
        "En la Ciudad de Quito el día, "
        f"{datos['fecha']} , dando cumplimiento "
        "a las disposiciones en la normativa "
        "legal vigente, se deja constancia que "
        "el estudiante"
    )

    y = parrafo_centrado(
        pdf,
        introduccion,
        18,
        y,
        ancho_pagina - 36,
        tamano=10.5,
        interlineado=13,
    )

    pdf.setFont(
        "Times-Bold",
        10.5,
    )

    pdf.drawCentredString(
        ancho_pagina / 2,
        y,
        datos["estudiante"],
    )

    y -= 14

    linea_estudiante = (
        f"{datos['cedula']}                 "
        f"{datos['presentacion']}"
    )

    y = parrafo_centrado(
        pdf,
        linea_estudiante,
        18,
        y,
        ancho_pagina - 36,
        tamano=10.5,
        interlineado=13,
    )

    y -= 36

    deliberacion = (
        "Luego de la deliberación del tribunal "
        "el estudiante ha obtenido como resultado "
        "las siguientes calificaciones, en la "
        "materia de Integración Curricular"
    )

    y = parrafo_izquierda(
        pdf,
        deliberacion,
        18,
        y,
        ancho_pagina - 36,
        tamano=10.5,
        interlineado=13,
    )

    y -= 18

    # Tabla de calificaciones

    if datos["complexivo"]:
        tabla_x = 18
        tabla_ancho = 350
        columna_1 = 190
        alturas = [
            34,
            34,
            20,
        ]

    else:
        tabla_x = 18
        tabla_ancho = (
            ancho_pagina - 36
        )
        columna_1 = (
            tabla_ancho * 0.72
        )
        alturas = [
            20,
            20,
            20,
        ]

    columna_2 = (
        tabla_ancho - columna_1
    )

    for (
        etiqueta,
        valor,
    ), altura in zip(
        datos["calificaciones"],
        alturas,
    ):
        pdf.rect(
            tabla_x,
            y - altura,
            columna_1,
            altura,
        )

        pdf.rect(
            tabla_x + columna_1,
            y - altura,
            columna_2,
            altura,
        )

        lineas = dividir_lineas(
            etiqueta,
            "Times-Bold",
            10,
            columna_1 - 8,
        )

        posicion_texto = (
            y - 12
            if len(lineas) == 1
            else y - 11
        )

        pdf.setFont(
            "Times-Bold",
            10,
        )

        for linea in lineas:
            pdf.drawString(
                tabla_x + 4,
                posicion_texto,
                linea,
            )

            posicion_texto -= 11

        pdf.setFont(
            "Times-Roman",
            10.5,
        )

        pdf.drawCentredString(
            tabla_x
            + columna_1
            + columna_2 / 2,
            y - altura / 2 - 3,
            valor,
        )

        y -= altura

    # Aviso institucional

    y -= 13

    aviso = (
        "Se informa que esta nota y este evento "
        "no constituyen una ceremonia de graduación. "
        "De acuerdo con la normativa vigente de la "
        "PUCE, la Secretaría certificará el cumplimiento "
        "de todos los requisitos y notificará al "
        "estudiante correspondiente."
    )

    parrafo_centrado(
        pdf,
        aviso,
        20,
        y,
        ancho_pagina - 40,
        fuente="Times-Roman",
        tamano=9.5,
        interlineado=12,
    )

    # Firmas

    if datos["complexivo"]:
        posiciones = [
            (18, 265, False),
            (
                ancho_pagina - 178,
                265,
                False,
            ),
            (18, 115, True),
        ]

    else:
        posiciones = [
            (18, 275, True),
            (
                ancho_pagina - 178,
                275,
                True,
            ),
            (18, 120, True),
            (
                ancho_pagina - 178,
                120,
                True,
            ),
        ]

    for posicion, nombre in zip(
        posiciones,
        datos["firmas"],
    ):
        dibujar_firma(
            pdf,
            posicion[0],
            posicion[1],
            nombre,
            posicion[2],
        )

    pdf.showPage()
    pdf.save()

    memoria.seek(0)

    return memoria.getvalue()


# ==========================================================
# WORD
# ==========================================================

def margenes_celda(
    celda,
    top=0,
    start=0,
    bottom=0,
    end=0,
):
    propiedades = (
        celda._tc.get_or_add_tcPr()
    )

    margenes = (
        propiedades.first_child_found_in(
            "w:tcMar"
        )
    )

    if margenes is None:
        margenes = OxmlElement(
            "w:tcMar"
        )

        propiedades.append(
            margenes
        )

    for nombre, valor in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        nodo = margenes.find(
            qn(f"w:{nombre}")
        )

        if nodo is None:
            nodo = OxmlElement(
                f"w:{nombre}"
            )

            margenes.append(
                nodo
            )

        nodo.set(
            qn("w:w"),
            str(valor),
        )

        nodo.set(
            qn("w:type"),
            "dxa",
        )


def bordes_tabla(
    tabla,
    tamano="12",
):
    propiedades = (
        tabla._tbl.tblPr
    )

    bordes = OxmlElement(
        "w:tblBorders"
    )

    for lado in [
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ]:
        elemento = OxmlElement(
            f"w:{lado}"
        )

        elemento.set(
            qn("w:val"),
            "single",
        )

        elemento.set(
            qn("w:sz"),
            tamano,
        )

        elemento.set(
            qn("w:color"),
            "444444",
        )

        bordes.append(
            elemento
        )

    propiedades.append(
        bordes
    )


def quitar_bordes(
    tabla,
):
    propiedades = (
        tabla._tbl.tblPr
    )

    bordes = OxmlElement(
        "w:tblBorders"
    )

    for lado in [
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ]:
        elemento = OxmlElement(
            f"w:{lado}"
        )

        elemento.set(
            qn("w:val"),
            "nil",
        )

        bordes.append(
            elemento
        )

    propiedades.append(
        bordes
    )


def generar_word_oficial(
    acta,
    logo_path,
):
    datos = contexto_acta(
        acta
    )

    documento = Document()

    seccion = documento.sections[0]

    seccion.page_width = Cm(21)
    seccion.page_height = Cm(29.7)

    seccion.top_margin = Cm(1.2)
    seccion.bottom_margin = Cm(1.2)
    seccion.left_margin = Cm(1.3)
    seccion.right_margin = Cm(1.3)

    estilo = documento.styles[
        "Normal"
    ]

    estilo.font.name = (
        "Times New Roman"
    )

    estilo.font.size = Pt(10)

    # Encabezado

    cabecera = documento.add_table(
        rows=1,
        cols=2,
    )

    cabecera.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    cabecera.autofit = False

    izquierda = cabecera.cell(
        0,
        0,
    )

    derecha = cabecera.cell(
        0,
        1,
    )

    izquierda.width = Cm(8.2)
    derecha.width = Cm(9.5)

    margenes_celda(
        izquierda
    )

    margenes_celda(
        derecha
    )

    p = izquierda.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    logo_path = Path(
        logo_path
    )

    if logo_path.exists():
        p.add_run().add_picture(
            str(logo_path),
            width=Cm(6.0),
        )

    p = derecha.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    run = p.add_run(
        "UNIDAD ACADÉMICA ESPECIALIZADA EN\n"
        "FORMACIÓN TÉCNICA Y TECNOLÓGICA\n"
        "PUCE TEC"
    )

    run.font.name = "Arial"
    run.font.size = Pt(8)

    # Título

    titulo = documento.add_paragraph()

    titulo.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    titulo.paragraph_format.space_before = (
        Pt(8)
    )

    titulo.paragraph_format.space_after = (
        Pt(38)
    )

    run = titulo.add_run(
        datos["titulo"]
    )

    run.bold = True
    run.font.name = (
        "Times New Roman"
    )

    run.font.size = Pt(12)

    # Introducción

    p = documento.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = (
        Pt(0)
    )

    p.add_run(
        "En la Ciudad de Quito el día, "
        f"{datos['fecha']} , dando cumplimiento "
        "a las\n"
        "disposiciones en la normativa legal "
        "vigente, se deja constancia que el "
        "estudiante"
    )

    # Estudiante

    p = documento.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = (
        Pt(0)
    )

    run = p.add_run(
        datos["estudiante"]
    )

    run.bold = True

    # Cédula + presentación

    p = documento.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = (
        Pt(28)
    )

    p.add_run(
        f"{datos['cedula']}                 "
        f"{datos['presentacion']}"
    )

    # Deliberación

    p = documento.add_paragraph(
        "Luego de la deliberación del tribunal "
        "el estudiante ha obtenido como resultado "
        "las siguientes\n"
        "calificaciones, en la materia de "
        "Integración Curricular"
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    p.paragraph_format.space_after = (
        Pt(10)
    )

    # Tabla

    tabla = documento.add_table(
        rows=0,
        cols=2,
    )

    tabla.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    tabla.autofit = False

    bordes_tabla(
        tabla
    )

    if datos["complexivo"]:
        ancho_1 = Cm(5.4)
        ancho_2 = Cm(4.5)

    else:
        ancho_1 = Cm(12.4)
        ancho_2 = Cm(4.8)

    for (
        etiqueta,
        valor,
    ) in datos["calificaciones"]:
        celdas = tabla.add_row().cells

        celdas[0].width = ancho_1
        celdas[1].width = ancho_2

        for celda in celdas:
            celda.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            margenes_celda(
                celda,
                start=40,
                end=40,
            )

        p1 = celdas[0].paragraphs[0]

        p1.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = p1.add_run(
            etiqueta
        )

        run.bold = True
        run.font.name = (
            "Times New Roman"
        )

        run.font.size = Pt(10)

        p2 = celdas[1].paragraphs[0]

        p2.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = p2.add_run(
            valor
        )

        run.font.name = (
            "Times New Roman"
        )

        run.font.size = Pt(10)

    # Aviso

    aviso = documento.add_paragraph(
        "Se informa que esta nota y este evento "
        "no constituyen una ceremonia de graduación. "
        "De acuerdo con\n"
        "la normativa vigente de la PUCE, la Secretaría "
        "certificará el cumplimiento de todos los "
        "requisitos y\n"
        "notificará al estudiante correspondiente."
    )

    aviso.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    aviso.paragraph_format.space_before = (
        Pt(0)
    )

    aviso.paragraph_format.space_after = (
        Pt(42)
    )

    # Firmas

    firmas = documento.add_table(
        rows=2,
        cols=2,
    )

    firmas.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    firmas.autofit = False

    quitar_bordes(
        firmas
    )

    indice = 0

    for fila in range(2):
        for columna in range(2):
            celda = firmas.cell(
                fila,
                columna,
            )

            margenes_celda(
                celda,
                start=200,
                end=200,
            )

            p = celda.paragraphs[0]

            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            if (
                indice
                < len(datos["firmas"])
            ):
                if fila == 1:
                    p.paragraph_format.space_before = (
                        Pt(45)
                    )

                p.add_run(
                    "____________________________\n"
                )

                mostrar_firma = (
                    not datos["complexivo"]
                    or indice == 2
                )

                if mostrar_firma:
                    run = p.add_run(
                        "Firma\n"
                    )

                    run.font.size = Pt(9)

                run = p.add_run(
                    datos["firmas"][
                        indice
                    ]
                )

                run.font.name = (
                    "Times New Roman"
                )

                run.font.size = Pt(9)

            indice += 1

    memoria = BytesIO()

    documento.save(
        memoria
    )

    memoria.seek(0)

    return memoria.getvalue()
